from fastapi import APIRouter, UploadFile, File, HTTPException, status, Depends
from sqlalchemy.orm import Session
from ..core.database import get_db
from ..models.document import Document
from ..core.config import settings
import os
import shutil
from datetime import datetime
from PyPDF2 import PdfReader
import docx

router = APIRouter()

UPLOAD_DIR = settings.UPLOAD_DIR if hasattr(settings, 'UPLOAD_DIR') else './uploads'
# DO NOT call os.makedirs here!

# Helper to extract text from PDF
def extract_pdf_text(file_path):
    try:
        reader = PdfReader(file_path)
        return '\n'.join(page.extract_text() or '' for page in reader.pages)
    except Exception:
        return ''

# Helper to extract text from DOCX
def extract_docx_text(file_path):
    try:
        doc = docx.Document(file_path)
        return '\n'.join([p.text for p in doc.paragraphs])
    except Exception:
        return ''

# Helper to extract text from TXT
def extract_txt_text(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception:
        return ''

@router.post('/documents/upload')
async def upload_document(file: UploadFile = File(...), db: Session = Depends(get_db)):
    # Ensure upload directory exists at runtime (not at import time)
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    # Save file to disk
    filename = f"{datetime.utcnow().strftime('%Y%m%d%H%M%S')}_{file.filename}"
    file_path = os.path.join(UPLOAD_DIR, filename)
    with open(file_path, 'wb') as out_file:
        shutil.copyfileobj(file.file, out_file)

    # Extract text
    ext = file.filename.lower().split('.')[-1]
    if ext == 'pdf':
        extracted_text = extract_pdf_text(file_path)
    elif ext in ('docx', 'doc'):
        extracted_text = extract_docx_text(file_path)
    elif ext == 'txt':
        extracted_text = extract_txt_text(file_path)
    else:
        extracted_text = ''

    # Store metadata and text in DB
    doc = Document(
        filename=filename,
        original_filename=file.filename,
        file_path=file_path,
        file_size=os.path.getsize(file_path),
        content_type=file.content_type,
        extracted_text=extracted_text,
        uploaded_at=datetime.utcnow(),
        status='processed',
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    return {"success": True, "document_id": doc.id, "filename": filename} 